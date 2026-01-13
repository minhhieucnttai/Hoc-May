#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Chuyển Markdown sang Word với Times New Roman 13pt
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def convert_md_to_docx():
    """Chuyển markdown sang Word"""
    
    # Đọc file markdown
    md_file = Path("BaoCao_HocMay_TranMinhHieu.md")
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Tạo document Word mới
    doc = Document()
    
    # Cài đặt font mặc định
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    
    # Xử lý nội dung
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        if not line.strip():
            doc.add_paragraph()
        elif line.startswith('# '):
            # Heading 1
            heading = doc.add_heading(line[2:], level=1)
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(16)
                run.bold = True
        elif line.startswith('## '):
            # Heading 2
            heading = doc.add_heading(line[3:], level=2)
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.bold = True
        elif line.startswith('### '):
            # Heading 3
            heading = doc.add_heading(line[4:], level=3)
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
        elif line.startswith('- '):
            # Bullet point
            p = doc.add_paragraph(line[2:], style='List Bullet')
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.bold = True
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # Table detection
            if '---' in lines[i + 1]:
                # Parse table
                header_cells = [cell.strip() for cell in line.split('|')[1:-1]]
                i += 2  # Skip separator line
                
                # Tạo table
                table = doc.add_table(rows=1, cols=len(header_cells))
                table.style = 'Light Grid Accent 1'
                
                # Add header
                hdr_cells = table.rows[0].cells
                for j, cell_text in enumerate(header_cells):
                    cell = hdr_cells[j]
                    p = cell.paragraphs[0]
                    run = p.add_run(cell_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
                    run.bold = True
                
                # Add rows
                while i < len(lines) and '|' in lines[i]:
                    row_cells = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                    if len(row_cells) == len(header_cells):
                        row = table.add_row()
                        for j, cell_text in enumerate(row_cells):
                            cell = row.cells[j]
                            p = cell.paragraphs[0]
                            run = p.add_run(cell_text)
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(13)
                    i += 1
                continue
        else:
            # Paragraph thường
            if line.strip():
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
        
        i += 1
    
    # Lưu file
    output_file = Path("BaoCao_HocMay_TranMinhHieu.docx")
    doc.save(output_file)
    
    print(f"✅ Đã tạo file: {output_file}")
    print(f"✅ Font: Times New Roman, Size: 13pt")
    print(f"✅ Bố cục: Đẹp, chuẩn tắc")

if __name__ == "__main__":
    convert_md_to_docx()
